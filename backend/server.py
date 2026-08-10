from fastapi import FastAPI, APIRouter, HTTPException, Depends, Header
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os, uuid, logging, bcrypt, jwt, io, re
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional
from datetime import datetime, timezone, timedelta

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

MONGO_URL = os.environ['MONGO_URL']
MONGO_DB = os.environ['MONGO_DB']
JWT_SECRET = os.environ['JWT_SECRET']
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')

client = AsyncIOMotorClient(MONGO_URL)
db = client[MONGO_DB]

app = FastAPI(title="FreshMart API")
api = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ---------- Utils ----------
def now_iso(): return datetime.now(timezone.utc).isoformat()

def hash_pw(p: str) -> str:
    return bcrypt.hashpw(p.encode(), bcrypt.gensalt()).decode()

def verify_pw(p: str, h: str) -> bool:
    try: return bcrypt.checkpw(p.encode(), h.encode())
    except: return False

def make_token(uid: str, role: str) -> str:
    payload = {"sub": uid, "role": role, "exp": datetime.now(timezone.utc) + timedelta(days=30)}
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")

async def get_current_user(authorization: Optional[str] = Header(None)):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(401, "Missing token")
    token = authorization.split(" ", 1)[1]
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
    except jwt.PyJWTError:
        raise HTTPException(401, "Invalid token")
    user = await db.users.find_one({"id": payload["sub"]}, {"_id": 0, "password_hash": 0})
    if not user: raise HTTPException(401, "User not found")
    return user

async def require_admin(user=Depends(get_current_user)):
    if user.get("role") != "admin":
        raise HTTPException(403, "Admin only")
    return user

# ---------- Models ----------
class RegisterIn(BaseModel):
    name: str
    email: EmailStr
    password: str

class LoginIn(BaseModel):
    email: EmailStr
    password: str

class ProductIn(BaseModel):
    name: str
    brand: Optional[str] = ""
    category: str
    price: float
    mrp: float
    weight: Optional[str] = ""
    image: str
    stock: int = 50
    rating: float = 4.5
    reviews: int = 0
    delivery_min: int = 15
    description: Optional[str] = ""
    tags: List[str] = []
    featured: bool = False

class CartItemIn(BaseModel):
    product_id: str
    qty: int = 1

class OrderIn(BaseModel):
    address: dict
    payment_method: str = "COD"
    coupon: Optional[str] = None

class ChatIn(BaseModel):
    message: str
    session_id: Optional[str] = None

class ProductUpdateIn(BaseModel):
    name: Optional[str] = None
    brand: Optional[str] = None
    category: Optional[str] = None
    price: Optional[float] = None
    mrp: Optional[float] = None
    weight: Optional[str] = None
    image: Optional[str] = None
    stock: Optional[int] = None
    rating: Optional[float] = None
    reviews: Optional[int] = None
    delivery_min: Optional[int] = None
    description: Optional[str] = None
    tags: Optional[List[str]] = None
    featured: Optional[bool] = None

class UserUpdateIn(BaseModel):
    name: Optional[str] = None
    role: Optional[str] = None
    wallet: Optional[float] = None

class OrderStatusIn(BaseModel):
    status: str

# ---------- Auth ----------
@api.post("/auth/register")
async def register(inp: RegisterIn):
    existing = await db.users.find_one({"email": inp.email.lower()})
    if existing:
        raise HTTPException(400, "Email already registered")
    uid = str(uuid.uuid4())
    role = "admin" if inp.email.lower() == "admin@freshmart.com" else "customer"
    user = {
        "id": uid, "name": inp.name, "email": inp.email.lower(),
        "password_hash": hash_pw(inp.password), "role": role,
        "created_at": now_iso(), "wallet": 100.0
    }
    await db.users.insert_one(user)
    return {"token": make_token(uid, role), "user": {"id": uid, "name": inp.name, "email": inp.email.lower(), "role": role, "wallet": 100.0}}

@api.post("/auth/login")
async def login(inp: LoginIn):
    user = await db.users.find_one({"email": inp.email.lower()})
    if not user or not verify_pw(inp.password, user["password_hash"]):
        raise HTTPException(401, "Invalid credentials")
    return {"token": make_token(user["id"], user["role"]), "user": {"id": user["id"], "name": user["name"], "email": user["email"], "role": user["role"], "wallet": user.get("wallet", 0)}}

@api.get("/auth/me")
async def me(user=Depends(get_current_user)):
    return user

# ---------- Categories ----------
CATEGORIES = [
    {"slug": "fruits", "name": "Fruits", "image": "https://images.unsplash.com/photo-1610348725531-843dff563e2c?w=400", "icon": "🍎"},
    {"slug": "vegetables", "name": "Vegetables", "image": "https://images.unsplash.com/photo-1590779033100-9f60a05a013d?w=400", "icon": "🥦"},
    {"slug": "dairy", "name": "Milk & Dairy", "image": "https://images.unsplash.com/photo-1679458219939-4f6347c55460?w=400", "icon": "🥛"},
    {"slug": "bakery", "name": "Bread & Bakery", "image": "https://images.unsplash.com/photo-1608198093002-ad4e005484ec?w=400", "icon": "🍞"},
    {"slug": "snacks", "name": "Snacks", "image": "https://images.unsplash.com/photo-1530590543196-9faf82fed4ea?w=400", "icon": "🍿"},
    {"slug": "beverages", "name": "Beverages", "image": "https://images.unsplash.com/photo-1625772299848-391b6a87d7b3?w=400", "icon": "🥤"},
    {"slug": "staples", "name": "Rice, Flour & Pulses", "image": "https://images.unsplash.com/photo-1586201375761-83865001e31c?w=400", "icon": "🌾"},
    {"slug": "personal-care", "name": "Personal Care", "image": "https://images.unsplash.com/photo-1556228720-195a672e8a03?w=400", "icon": "🧴"},
    {"slug": "cleaning", "name": "Cleaning Supplies", "image": "https://images.unsplash.com/photo-1585421514738-01798e348b17?w=400", "icon": "🧽"},
    {"slug": "kitchen", "name": "Kitchen Items", "image": "https://images.unsplash.com/photo-1556909114-f6e7ad7d3136?w=400", "icon": "🍳"},
]

@api.get("/categories")
async def get_categories():
    return CATEGORIES

# ---------- Products ----------
@api.get("/products")
async def list_products(category: Optional[str] = None, q: Optional[str] = None, featured: Optional[bool] = None, limit: int = 100):
    query = {}
    if category: query["category"] = category
    if featured is not None: query["featured"] = featured
    if q:
        safe_q = re.escape(q.strip())
        query["$or"] = [{"name": {"$regex": safe_q, "$options": "i"}}, {"brand": {"$regex": safe_q, "$options": "i"}}, {"tags": {"$regex": safe_q, "$options": "i"}}]
    items = await db.products.find(query, {"_id": 0}).limit(limit).to_list(limit)
    return items

@api.get("/products/{pid}")
async def get_product(pid: str):
    p = await db.products.find_one({"id": pid}, {"_id": 0})
    if not p: raise HTTPException(404, "Not found")
    similar = await db.products.find({"category": p["category"], "id": {"$ne": pid}}, {"_id": 0}).limit(6).to_list(6)
    return {"product": p, "similar": similar}

@api.post("/products")
async def create_product(inp: ProductIn, admin=Depends(require_admin)):
    pid = str(uuid.uuid4())
    doc = inp.model_dump()
    discount = round((inp.mrp - inp.price) / inp.mrp * 100) if inp.mrp > inp.price else 0
    doc.update({"id": pid, "discount": discount, "created_at": now_iso()})
    await db.products.insert_one(doc)
    doc.pop("_id", None)
    return doc

@api.put("/products/{pid}")
async def update_product(pid: str, inp: ProductUpdateIn, admin=Depends(require_admin)):
    existing = await db.products.find_one({"id": pid}, {"_id": 0})
    if not existing:
        raise HTTPException(404, "Not found")

    updates = {k: v for k, v in inp.model_dump().items() if v is not None}
    if updates:
        new_price = updates.get("price", existing["price"])
        new_mrp = updates.get("mrp", existing["mrp"])
        if new_mrp > new_price:
            updates["discount"] = round((new_mrp - new_price) / new_mrp * 100)
        await db.products.update_one({"id": pid}, {"$set": updates})

    was_out_of_stock = existing.get("stock", 0) <= 0
    is_now_in_stock = updates.get("stock", existing.get("stock", 0)) > 0

    if was_out_of_stock and is_now_in_stock:
        await notify_back_in_stock(pid)

    doc = await db.products.find_one({"id": pid}, {"_id": 0})
    return doc

@api.delete("/products/{pid}")
async def delete_product(pid: str, admin=Depends(require_admin)):
    await db.products.delete_one({"id": pid})
    return {"ok": True}

# ---------- Back-in-stock notifications ----------
async def notify_back_in_stock(pid: str):
    product = await db.products.find_one({"id": pid}, {"_id": 0})
    if not product:
        return
    subs = await db.notify_requests.find({"product_id": pid}, {"_id": 0}).to_list(1000)
    if not subs:
        return
    notifications = [{
        "id": str(uuid.uuid4()),
        "user_id": s["user_id"],
        "type": "back_in_stock",
        "product_id": pid,
        "product_name": product["name"],
        "message": f'"{product["name"]}" is back in stock!',
        "read": False,
        "created_at": now_iso(),
    } for s in subs]
    await db.notifications.insert_many(notifications)
    await db.notify_requests.delete_many({"product_id": pid})

@api.post("/products/{pid}/notify-me")
async def notify_me(pid: str, user=Depends(get_current_user)):
    product = await db.products.find_one({"id": pid}, {"_id": 0})
    if not product:
        raise HTTPException(404, "Not found")
    existing = await db.notify_requests.find_one({"product_id": pid, "user_id": user["id"]})
    if existing:
        return {"ok": True, "message": "You're already on the notify list for this product"}
    await db.notify_requests.insert_one({
        "id": str(uuid.uuid4()), "product_id": pid, "user_id": user["id"], "created_at": now_iso()
    })
    return {"ok": True, "message": "We'll notify you when this is back in stock"}

@api.get("/notifications")
async def list_notifications(user=Depends(get_current_user)):
    items = await db.notifications.find({"user_id": user["id"]}, {"_id": 0}).sort("created_at", -1).to_list(100)
    unread_count = sum(1 for n in items if not n.get("read"))
    return {"notifications": items, "unread_count": unread_count}

@api.put("/notifications/{nid}/read")
async def mark_notification_read(nid: str, user=Depends(get_current_user)):
    await db.notifications.update_one({"id": nid, "user_id": user["id"]}, {"$set": {"read": True}})
    return {"ok": True}

@api.put("/notifications/read-all")
async def mark_all_notifications_read(user=Depends(get_current_user)):
    await db.notifications.update_many({"user_id": user["id"], "read": False}, {"$set": {"read": True}})
    return {"ok": True}

# ---------- Cart ----------
@api.get("/cart")
async def get_cart(user=Depends(get_current_user)):
    cart = await db.carts.find_one({"user_id": user["id"]}, {"_id": 0})
    if not cart: return {"items": [], "products": []}
    pids = [i["product_id"] for i in cart.get("items", [])]
    products = await db.products.find({"id": {"$in": pids}}, {"_id": 0}).to_list(100)
    return {"items": cart.get("items", []), "products": products}

@api.post("/cart")
async def add_cart(inp: CartItemIn, user=Depends(get_current_user)):
    cart = await db.carts.find_one({"user_id": user["id"]})
    items = cart.get("items", []) if cart else []
    found = next((i for i in items if i["product_id"] == inp.product_id), None)
    if found:
        found["qty"] = max(1, found["qty"] + inp.qty) if inp.qty != 0 else found["qty"]
        if inp.qty < 0 and found["qty"] <= 0:
            items = [i for i in items if i["product_id"] != inp.product_id]
    else:
        if inp.qty > 0:
            items.append({"product_id": inp.product_id, "qty": inp.qty})
    await db.carts.update_one({"user_id": user["id"]}, {"$set": {"user_id": user["id"], "items": items, "updated_at": now_iso()}}, upsert=True)
    return {"items": items}

@api.put("/cart/{pid}")
async def set_cart_qty(pid: str, qty: int, user=Depends(get_current_user)):
    cart = await db.carts.find_one({"user_id": user["id"]})
    items = cart.get("items", []) if cart else []
    items = [i for i in items if i["product_id"] != pid]
    if qty > 0:
        items.append({"product_id": pid, "qty": qty})
    await db.carts.update_one({"user_id": user["id"]}, {"$set": {"items": items}}, upsert=True)
    return {"items": items}

@api.delete("/cart/{pid}")
async def del_cart(pid: str, user=Depends(get_current_user)):
    await db.carts.update_one({"user_id": user["id"]}, {"$pull": {"items": {"product_id": pid}}})
    return {"ok": True}

# ---------- Orders ----------
@api.post("/orders")
async def create_order(inp: OrderIn, user=Depends(get_current_user)):
    cart = await db.carts.find_one({"user_id": user["id"]})
    if not cart or not cart.get("items"):
        raise HTTPException(400, "Cart is empty")
    pids = [i["product_id"] for i in cart["items"]]
    products = await db.products.find({"id": {"$in": pids}}, {"_id": 0}).to_list(100)
    pmap = {p["id"]: p for p in products}
    order_items = []
    subtotal = 0.0
    for i in cart["items"]:
        p = pmap.get(i["product_id"])
        if not p: continue
        order_items.append({"product_id": p["id"], "name": p["name"], "image": p["image"], "price": p["price"], "qty": i["qty"]})
        subtotal += p["price"] * i["qty"]
    discount = 0
    if inp.coupon and inp.coupon.upper() == "FRESH10":
        discount = round(subtotal * 0.10, 2)
    delivery_fee = 0 if subtotal >= 199 else 29
    total = round(subtotal - discount + delivery_fee, 2)
    oid = str(uuid.uuid4())
    order = {
        "id": oid, "user_id": user["id"], "items": order_items, "address": inp.address,
        "payment_method": inp.payment_method, "coupon": inp.coupon,
        "subtotal": round(subtotal, 2), "discount": discount, "delivery_fee": delivery_fee, "total": total,
        "status": "confirmed", "created_at": now_iso(),
        "eta_minutes": 30
    }
    await db.orders.insert_one(order)
    await db.carts.update_one({"user_id": user["id"]}, {"$set": {"items": []}})
    order.pop("_id", None)
    return order

@api.get("/orders")
async def list_orders(user=Depends(get_current_user)):
    orders = await db.orders.find({"user_id": user["id"]}, {"_id": 0}).sort("created_at", -1).to_list(100)
    return orders

@api.get("/orders/{oid}")
async def get_order(oid: str, user=Depends(get_current_user)):
    o = await db.orders.find_one({"id": oid, "user_id": user["id"]}, {"_id": 0})
    if not o: raise HTTPException(404, "Not found")
    return o

# ---------- Invoices (PDF / Word) ----------
def _build_invoice_pdf(order: dict, user: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20 * mm, bottomMargin=20 * mm)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("FreshMart", styles["Title"]))
    elements.append(Paragraph("Order Invoice", styles["Heading2"]))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(f"Order ID: {order['id']}", styles["Normal"]))
    elements.append(Paragraph(f"Date: {order['created_at']}", styles["Normal"]))
    elements.append(Paragraph(f"Status: {order['status']}", styles["Normal"]))
    elements.append(Paragraph(f"Customer: {user.get('name', '')} ({user.get('email', '')})", styles["Normal"]))
    addr = order.get("address", {})
    addr_line = ", ".join(str(v) for v in [addr.get("line1"), addr.get("city"), addr.get("pincode")] if v)
    elements.append(Paragraph(f"Delivery address: {addr_line}", styles["Normal"]))
    elements.append(Paragraph(f"Payment method: {order.get('payment_method', 'COD')}", styles["Normal"]))
    elements.append(Spacer(1, 12))

    table_data = [["Product", "Price", "Qty", "Subtotal"]]
    for item in order["items"]:
        table_data.append([
            item["name"], f"₹{item['price']}", str(item["qty"]), f"₹{item['price'] * item['qty']:.2f}"
        ])
    table = Table(table_data, colWidths=[80 * mm, 30 * mm, 20 * mm, 30 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#22c55e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("ALIGN", (0, 0), (0, -1), "LEFT"),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 12))

    summary_data = [
        ["Subtotal", f"₹{order['subtotal']:.2f}"],
        ["Discount", f"-₹{order.get('discount', 0):.2f}"],
        ["Delivery fee", f"₹{order.get('delivery_fee', 0):.2f}"],
        ["Total", f"₹{order['total']:.2f}"],
    ]
    summary = Table(summary_data, colWidths=[130 * mm, 30 * mm])
    summary.setStyle(TableStyle([
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("LINEABOVE", (0, -1), (-1, -1), 1, colors.black),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
    ]))
    elements.append(summary)
    elements.append(Spacer(1, 20))
    elements.append(Paragraph("Thank you for shopping with FreshMart!", styles["Normal"]))

    doc.build(elements)
    return buf.getvalue()

def _build_invoice_docx(order: dict, user: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.add_heading("FreshMart", level=0)
    document.add_heading("Order Invoice", level=2)

    document.add_paragraph(f"Order ID: {order['id']}")
    document.add_paragraph(f"Date: {order['created_at']}")
    document.add_paragraph(f"Status: {order['status']}")
    document.add_paragraph(f"Customer: {user.get('name', '')} ({user.get('email', '')})")
    addr = order.get("address", {})
    addr_line = ", ".join(str(v) for v in [addr.get("line1"), addr.get("city"), addr.get("pincode")] if v)
    document.add_paragraph(f"Delivery address: {addr_line}")
    document.add_paragraph(f"Payment method: {order.get('payment_method', 'COD')}")

    table = document.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Product", "Price", "Qty", "Subtotal"
    for item in order["items"]:
        row = table.add_row().cells
        row[0].text = item["name"]
        row[1].text = f"₹{item['price']}"
        row[2].text = str(item["qty"])
        row[3].text = f"₹{item['price'] * item['qty']:.2f}"

    document.add_paragraph("")
    document.add_paragraph(f"Subtotal: ₹{order['subtotal']:.2f}")
    document.add_paragraph(f"Discount: -₹{order.get('discount', 0):.2f}")
    document.add_paragraph(f"Delivery fee: ₹{order.get('delivery_fee', 0):.2f}")
    total_p = document.add_paragraph()
    total_run = total_p.add_run(f"Total: ₹{order['total']:.2f}")
    total_run.bold = True
    total_run.font.size = Pt(13)

    document.add_paragraph("")
    document.add_paragraph("Thank you for shopping with FreshMart!")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()

@api.get("/orders/{oid}/invoice")
async def download_invoice(oid: str, format: str = "pdf", user=Depends(get_current_user)):
    order = await db.orders.find_one({"id": oid, "user_id": user["id"]}, {"_id": 0})
    if not order:
        raise HTTPException(404, "Order not found")

    try:
        if format == "docx":
            content = _build_invoice_docx(order, user)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"freshmart-invoice-{oid[:8]}.docx"
        else:
            content = _build_invoice_pdf(order, user)
            media_type = "application/pdf"
            filename = f"freshmart-invoice-{oid[:8]}.pdf"
    except ImportError as e:
        raise HTTPException(
            500,
            "Invoice generation requires 'reportlab' (for PDF) and 'python-docx' (for Word) "
            "to be installed on the backend. Run: pip install reportlab python-docx"
        )

    return StreamingResponse(
        io.BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

# ---------- Admin ----------
@api.get("/admin/stats")
async def admin_stats(admin=Depends(require_admin)):
    users_count = await db.users.count_documents({})
    orders = await db.orders.find({}, {"_id": 0}).to_list(1000)
    products_count = await db.products.count_documents({})
    revenue = sum(o.get("total", 0) for o in orders)
    return {
        "users": users_count, "orders": len(orders),
        "products": products_count, "revenue": round(revenue, 2),
        "recent_orders": orders[:10]
    }

@api.get("/admin/orders")
async def admin_orders(admin=Depends(require_admin)):
    return await db.orders.find({}, {"_id": 0}).sort("created_at", -1).to_list(200)

@api.put("/admin/orders/{oid}")
async def admin_update_order(oid: str, inp: OrderStatusIn, admin=Depends(require_admin)):
    result = await db.orders.update_one({"id": oid}, {"$set": {"status": inp.status}})
    if result.matched_count == 0:
        raise HTTPException(404, "Order not found")
    return await db.orders.find_one({"id": oid}, {"_id": 0})

@api.delete("/admin/orders/{oid}")
async def admin_delete_order(oid: str, admin=Depends(require_admin)):
    await db.orders.delete_one({"id": oid})
    return {"ok": True}

# ---------- Admin: Users CRUD ----------
@api.get("/admin/users")
async def admin_list_users(admin=Depends(require_admin)):
    return await db.users.find({}, {"_id": 0, "password_hash": 0}).sort("created_at", -1).to_list(1000)

@api.get("/admin/users/{uid}")
async def admin_get_user(uid: str, admin=Depends(require_admin)):
    u = await db.users.find_one({"id": uid}, {"_id": 0, "password_hash": 0})
    if not u:
        raise HTTPException(404, "User not found")
    return u

@api.put("/admin/users/{uid}")
async def admin_update_user(uid: str, inp: UserUpdateIn, admin=Depends(require_admin)):
    updates = {k: v for k, v in inp.model_dump().items() if v is not None}
    if updates.get("role") not in (None, "admin", "customer"):
        raise HTTPException(400, "role must be 'admin' or 'customer'")
    if updates:
        result = await db.users.update_one({"id": uid}, {"$set": updates})
        if result.matched_count == 0:
            raise HTTPException(404, "User not found")
    return await db.users.find_one({"id": uid}, {"_id": 0, "password_hash": 0})

@api.delete("/admin/users/{uid}")
async def admin_delete_user(uid: str, admin=Depends(require_admin)):
    if uid == admin["id"]:
        raise HTTPException(400, "You cannot delete your own admin account")
    await db.users.delete_one({"id": uid})
    return {"ok": True}

# ---------- AI Assistant ----------
@api.post("/ai/chat")
async def ai_chat(inp: ChatIn):
    """Optional AI assistant.

    The app runs without an external AI package/key. If EMERGENT_LLM_KEY is
    configured and the optional integration is installed, it will stream AI
    responses; otherwise it returns a helpful local fallback.
    """
    if not EMERGENT_LLM_KEY:
        return StreamingResponse(
            iter([f"FreshMart tip: I can help with groceries, recipes, healthy alternatives, and product categories. You asked: {inp.message}"]),
            media_type="text/plain"
        )

    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage, TextDelta, StreamDone
    except ImportError:
        return StreamingResponse(
            iter(["The AI integration is not installed. The rest of FreshMart is ready to use."]),
            media_type="text/plain"
        )

    session_id = inp.session_id or str(uuid.uuid4())
    system = (
        "You are FreshMart's friendly AI shopping assistant. Help users find groceries, "
        "suggest recipes, recommend healthy alternatives, and answer questions about "
        "fruits, vegetables, dairy, snacks, beverages, cleaning supplies, and personal care. "
        "Keep responses concise (under 100 words), warm, and use bullet points for lists. "
        "If asked non-grocery questions, gently redirect to shopping topics."
    )
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=session_id,
        system_message=system
    ).with_model("anthropic", "claude-sonnet-4-5-20250929")

    async def stream():
        try:
            async for ev in chat.stream_message(UserMessage(text=inp.message)):
                if isinstance(ev, TextDelta):
                    yield ev.content
                elif isinstance(ev, StreamDone):
                    break
        except Exception:
            logger.exception("AI error")
            yield "\n\nSorry, our assistant is momentarily unavailable."

    return StreamingResponse(
        stream(),
        media_type="text/plain",
        headers={"X-Accel-Buffering": "no", "Cache-Control": "no-cache"}
    )

# ---------- Seed ----------
@api.post("/seed")
async def seed():
    if await db.products.count_documents({}) > 0:
        return {"seeded": False, "message": "Already seeded"}
    # Ensure admin user
    if not await db.users.find_one({"email": "admin@freshmart.com"}):
        await db.users.insert_one({
            "id": str(uuid.uuid4()), "name": "Admin", "email": "admin@freshmart.com",
            "password_hash": hash_pw("admin123"), "role": "admin", "created_at": now_iso(), "wallet": 0
        })
    # Ensure demo user
    if not await db.users.find_one({"email": "demo@freshmart.com"}):
        await db.users.insert_one({
            "id": str(uuid.uuid4()), "name": "Demo Shopper", "email": "demo@freshmart.com",
            "password_hash": hash_pw("demo123"), "role": "customer", "created_at": now_iso(), "wallet": 100.0
        })

    seed_data = [
        # Fruits
        ("Fresh Bananas", "Farm Fresh", "fruits", 49, 69, "1 kg", "https://images.unsplash.com/photo-1571771894821-ce9b6c11b08e?w=500", ["fruit", "banana"], True),
        ("Red Apples Shimla", "Kashmir Best", "fruits", 189, 249, "1 kg", "https://images.unsplash.com/photo-1568702846914-96b305d2aaeb?w=500", ["apple", "fruit"], True),
        ("Alphonso Mangoes", "Ratnagiri", "fruits", 599, 799, "1 kg", "https://images.unsplash.com/photo-1553279768-865429fa0078?w=500", ["mango"], False),
        ("Sweet Oranges", "Nagpur", "fruits", 89, 120, "1 kg", "https://images.unsplash.com/photo-1611080626919-7cf5a9dbab5b?w=500", ["orange"], False),
        ("Green Grapes", "Nashik", "fruits", 79, 99, "500 g", "https://images.unsplash.com/photo-1599819811279-d5ad9cccf838?w=500", ["grapes"], False),
        # Vegetables
        ("Fresh Tomatoes", "Local Farm", "vegetables", 39, 55, "1 kg", "https://images.unsplash.com/photo-1546470427-e5ac89cd0b31?w=500", ["tomato"], True),
        ("Onions", "Nashik Best", "vegetables", 35, 50, "1 kg", "https://images.unsplash.com/photo-1618512496248-a07fe83aa8cb?w=500", ["onion"], False),
        ("Potatoes", "UP Farm", "vegetables", 29, 40, "1 kg", "https://images.unsplash.com/photo-1518977676601-b53f82aba655?w=500", ["potato"], False),
        ("Broccoli", "Organic Valley", "vegetables", 79, 110, "500 g", "https://images.unsplash.com/photo-1459411552884-841db9b3cc2a?w=500", ["broccoli"], False),
        ("Baby Spinach", "Green Leaf", "vegetables", 45, 60, "200 g", "https://images.unsplash.com/photo-1576045057995-568f588f82fb?w=500", ["spinach"], False),
        # Dairy
        ("Full Cream Milk", "Amul", "dairy", 68, 72, "1 L", "https://images.unsplash.com/photo-1550583724-b2692b85b150?w=500", ["milk"], True),
        ("Greek Yogurt", "Epigamia", "dairy", 65, 85, "150 g", "https://images.unsplash.com/photo-1488477181946-6428a0291777?w=500", ["yogurt"], False),
        ("Cheddar Cheese", "Amul", "dairy", 220, 260, "200 g", "https://images.unsplash.com/photo-1486297678162-eb2a19b0a32d?w=500", ["cheese"], False),
        ("Butter", "Amul", "dairy", 55, 60, "100 g", "https://images.unsplash.com/photo-1589985270826-4b7bb135bc9d?w=500", ["butter"], False),
        # Bakery
        ("Whole Wheat Bread", "Britannia", "bakery", 45, 55, "400 g", "https://images.unsplash.com/photo-1509440159596-0249088772ff?w=500", ["bread"], True),
        ("Croissants", "Le Bakery", "bakery", 149, 199, "4 pcs", "https://images.unsplash.com/photo-1555507036-ab1f4038808a?w=500", ["croissant"], False),
        ("Chocolate Muffins", "Sweet Co", "bakery", 99, 129, "6 pcs", "https://images.unsplash.com/photo-1607958996333-41aef7caefaa?w=500", ["muffin"], False),
        # Snacks
        ("Lay's Classic Chips", "Lay's", "snacks", 20, 25, "52 g", "https://images.unsplash.com/photo-1613919113640-25732ec5e61f?w=500", ["chips"], True),
        ("Roasted Almonds", "Nutty Yogi", "snacks", 349, 449, "500 g", "https://images.unsplash.com/photo-1508061253366-f7da158b6d46?w=500", ["nuts", "almond"], True),
        ("Dark Chocolate", "Cadbury", "snacks", 129, 160, "80 g", "https://images.unsplash.com/photo-1606312619070-d48b4c652a52?w=500", ["chocolate"], False),
        ("Oreo Cookies", "Cadbury", "snacks", 30, 40, "120 g", "https://images.unsplash.com/photo-1558961363-fa8fdf82db35?w=500", ["cookies"], False),
        # Beverages
        ("Coca-Cola", "Coke", "beverages", 40, 45, "750 ml", "https://images.unsplash.com/photo-1554866585-cd94860890b7?w=500", ["cola", "drink"], True),
        ("Fresh Orange Juice", "Tropicana", "beverages", 99, 130, "1 L", "https://images.unsplash.com/photo-1621506289937-a8e4df240d0b?w=500", ["juice"], False),
        ("Green Tea", "Tetley", "beverages", 189, 250, "100 bags", "https://images.unsplash.com/photo-1594631252845-29fc4cc8cde9?w=500", ["tea"], False),
        ("Cold Brew Coffee", "Blue Tokai", "beverages", 249, 299, "250 ml", "https://images.unsplash.com/photo-1461023058943-07fcbe16d735?w=500", ["coffee"], False),
        # Staples
        ("Basmati Rice", "India Gate", "staples", 320, 400, "1 kg", "https://images.unsplash.com/photo-1586201375761-83865001e31c?w=500", ["rice"], True),
        ("Whole Wheat Atta", "Aashirvaad", "staples", 340, 380, "5 kg", "https://images.unsplash.com/photo-1509440159596-0249088772ff?w=500", ["flour"], False),
        ("Toor Dal", "Tata Sampann", "staples", 189, 220, "1 kg", "https://images.unsplash.com/photo-1585996429466-cd23787739d3?w=500", ["dal", "pulses"], False),
        # Personal care
        ("Dove Body Wash", "Dove", "personal-care", 199, 249, "250 ml", "https://images.unsplash.com/photo-1556228720-195a672e8a03?w=500", ["soap"], False),
        ("Colgate Toothpaste", "Colgate", "personal-care", 89, 110, "150 g", "https://images.unsplash.com/photo-1607613009820-a29f7bb81c04?w=500", ["toothpaste"], False),
        # Cleaning
        ("Vim Dish Wash", "Vim", "cleaning", 145, 175, "500 ml", "https://images.unsplash.com/photo-1585421514738-01798e348b17?w=500", ["dishwash"], False),
        ("Surf Excel Detergent", "Surf", "cleaning", 289, 345, "1 kg", "https://images.unsplash.com/photo-1604335398980-ededbadcf1c5?w=500", ["detergent"], False),
        # Kitchen
        ("Non-Stick Frying Pan", "Prestige", "kitchen", 899, 1499, "24 cm", "https://images.unsplash.com/photo-1584990347449-716e6a1ec4c9?w=500", ["pan"], False),
        ("Steel Water Bottle", "Milton", "kitchen", 349, 499, "1 L", "https://images.unsplash.com/photo-1602143407151-7111542de6e8?w=500", ["bottle"], False),
    ]
    products = []
    for name, brand, cat, price, mrp, weight, image, tags, featured in seed_data:
        discount = round((mrp - price) / mrp * 100)
        products.append({
            "id": str(uuid.uuid4()), "name": name, "brand": brand, "category": cat,
            "price": price, "mrp": mrp, "discount": discount, "weight": weight,
            "image": image, "stock": 50, "rating": round(4.0 + (hash(name) % 10) / 10, 1),
            "reviews": 20 + (hash(name) % 500), "delivery_min": 15 + (hash(name) % 20),
            "description": f"Premium quality {name.lower()}. Fresh and delivered fast to your doorstep.",
            "tags": tags, "featured": featured, "created_at": now_iso()
        })
    await db.products.insert_many(products)
    return {"seeded": True, "count": len(products)}

@api.get("/")
async def root():
    return {"service": "FreshMart API", "status": "ok"}

app.include_router(api)
app.add_middleware(
    CORSMiddleware, allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"], allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown():
    client.close()